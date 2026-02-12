#!/usr/bin/env python3
"""
Interview Transcription System

Structured interview transcription using faster-whisper with automatic
CUDA/CPU detection. Produces clean Q&A formatted output with timestamps,
speaker detection, and optional interview guide matching.
"""

import argparse
import json
import os
import re
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Optional

# ──────────────────────────────────────────────────────────────
# Configuration — Modify these defaults to match your setup.
# All values can also be overridden via CLI flags (--help).
# ──────────────────────────────────────────────────────────────

MODEL_SIZE      = "large-v3"    # tiny, base, small, medium, large-v2, large-v3
LANGUAGE        = "sv"          # ISO 639-1: sv, en, de, fr, es, etc.
BEAM_SIZE       = 5             # Higher = more accurate, slower (1-10)
PAUSE_THRESHOLD = 2.0           # Seconds of silence to detect speaker change

# Context prompt — adjust for your language/interview style
INITIAL_PROMPT = (
    "Det här är en intervju på svenska. "
    "Intervjuaren ställer frågor och respondenten svarar."
)

AUDIO_EXTENSIONS = {
    ".mp3", ".wav", ".m4a", ".flac", ".ogg",
    ".wma", ".aac", ".mp4", ".mkv", ".webm",
}


# ──────────────────────────────────────────────────────────────
# Dependency Check
# ──────────────────────────────────────────────────────────────

def check_dependencies():
    missing = []
    for module, package in [
        ("faster_whisper", "faster-whisper"),
        ("docx", "python-docx"),
        ("rich", "rich"),
    ]:
        try:
            __import__(module)
        except ImportError:
            missing.append(package)
    if missing:
        print(f"Missing: {', '.join(missing)}")
        print(f"Install: pip install {' '.join(missing)}")
        sys.exit(1)


# ──────────────────────────────────────────────────────────────
# Device Detection
# ──────────────────────────────────────────────────────────────

def detect_device(requested: str = "auto") -> tuple[str, str]:
    """
    Returns (device, compute_type). Prefers CUDA, falls back to CPU.
    """
    if requested == "cpu":
        return "cpu", "int8"

    cuda_available = False
    if requested in ("cuda", "auto"):
        try:
            import torch
            cuda_available = torch.cuda.is_available()
            if cuda_available:
                try:
                    gpu_name = torch.cuda.get_device_name(0)
                    props = torch.cuda.get_device_properties(0)
                    vram = props.total_memory / (1024 ** 3)
                    print(f"  GPU detected: {gpu_name} ({vram:.1f} GB)")

                    # Verify the GPU is actually usable by running a small tensor op
                    test = torch.zeros(1, device="cuda")
                    del test
                except Exception as e:
                    print(f"  GPU found but not usable: {e}")
                    cuda_available = False
        except ImportError:
            try:
                import ctranslate2
                cuda_available = "cuda" in ctranslate2.get_supported_compute_types("cuda")
            except Exception:
                pass

    if cuda_available:
        return "cuda", "float16"

    if requested == "cuda":
        print("WARNING: CUDA requested but unavailable — falling back to CPU.")
    elif requested == "auto":
        print("  CUDA unavailable — using CPU (int8)")

    return "cpu", "int8"


# ──────────────────────────────────────────────────────────────
# Data Structures
# ──────────────────────────────────────────────────────────────

@dataclass
class Config:
    model_size: str = MODEL_SIZE
    device: str = "auto"
    compute_type: str = "auto"
    language: str = LANGUAGE
    beam_size: int = BEAM_SIZE
    best_of: int = 5
    patience: float = 1.0
    vad_filter: bool = True
    vad_parameters: dict = field(default_factory=lambda: {
        "min_silence_duration_ms": 600,
        "speech_pad_ms": 200,
        "threshold": 0.35,
        "min_speech_duration_ms": 250,
    })
    merge_gap_threshold: float = 1.5
    min_segment_words: int = 3
    pause_for_turn_switch: float = PAUSE_THRESHOLD
    guide_path: Optional[str] = None
    output_format: str = "docx"


@dataclass
class Segment:
    start: float
    end: float
    text: str
    speaker: str = ""


# ──────────────────────────────────────────────────────────────
# Interview Guide Parser
# ──────────────────────────────────────────────────────────────

def parse_interview_guide(path: str) -> list[dict]:
    if not path or not os.path.exists(path):
        return []

    ext = Path(path).suffix.lower()

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        except Exception as e:
            print(f"WARNING: Could not read guide .docx: {e}")
            return []
    else:
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        content = None
        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if content is None:
            print(f"WARNING: Could not decode guide file: {path}")
            return []
        lines = [l.strip() for l in content.splitlines() if l.strip()]

    # Filter lines to extract actual interview questions.
    # A question is a line that:
    #   - Is a real sentence (> 5 words), AND
    #   - Contains a question mark, OR starts with a question word, OR starts with numbering
    # Skip: instructions, single-word sub-bullets, headers
    question_words = {
        "hur", "vad", "var", "vem", "vilken", "vilka", "vilket", "varför",
        "när", "kan", "skulle", "är", "finns", "berätta", "beskriv", "förklara",
        "om", "tycker", "anser", "upplever", "känner",
        "how", "what", "where", "who", "which", "why", "when", "can",
        "could", "would", "do", "does", "tell", "describe", "explain",
    }

    questions = []
    for line in lines:
        # Strip numbering prefix
        cleaned = re.sub(r'^\s*\d+[\.\):\s]+', '', line).strip()
        if not cleaned:
            continue

        words = cleaned.split()

        # Skip very short lines (sub-bullets like "Ålder", "typ av arbete")
        if len(words) < 5:
            continue

        # Skip instruction lines (typically the first line mentioning "introduktionen")
        if any(skip in cleaned.lower() for skip in ["introduktion", "nämn i ", "obs:", "note:"]):
            continue

        first_word = words[0].lower().rstrip(".,;:?!")

        is_question = (
            "?" in cleaned or
            first_word in question_words or
            cleaned.lower().startswith("vad innebär") or
            cleaned.lower().startswith("om du ")
        )

        # Short lines (< 8 words) must have a question mark or be a command-style
        # question (berätta, beskriv, förklara) to qualify
        command_starters = {"berätta", "beskriv", "förklara", "tell", "describe", "explain"}
        if len(words) < 8 and "?" not in cleaned and first_word not in command_starters:
            is_question = False

        if is_question:
            questions.append({"number": len(questions) + 1, "text": cleaned})

    if questions:
        print(f"  Interview guide: {len(questions)} questions loaded")
        for q in questions[:3]:
            print(f"    Q{q['number']}: {q['text'][:80]}...")
        if len(questions) > 3:
            print(f"    ... and {len(questions) - 3} more")

    return questions


# ──────────────────────────────────────────────────────────────
# Transcription Engine
# ──────────────────────────────────────────────────────────────

class Transcriber:
    def __init__(self, config: Config):
        self.config = config
        self.model = None
        self.guide_questions = parse_interview_guide(config.guide_path) if config.guide_path else []

    def load_model(self):
        from faster_whisper import WhisperModel
        from rich.console import Console
        console = Console()

        device, compute_type = detect_device(self.config.device)
        self.config.device = device
        self.config.compute_type = compute_type

        with console.status(f"[bold cyan]Loading {self.config.model_size} on {device} ({compute_type})..."):
            self.model = WhisperModel(
                self.config.model_size,
                device=device,
                compute_type=compute_type,
                download_root=os.path.expanduser("~/.cache/whisper-models"),
            )
        console.print(f"[green]✓[/green] Model: {self.config.model_size} ({device}, {compute_type})")

    def transcribe(self, audio_path: str) -> list[Segment]:
        from rich.console import Console
        from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
        console = Console()

        if self.model is None:
            self.load_model()

        console.print(f"\n[bold]Transcribing:[/bold] {os.path.basename(audio_path)}")
        start_time = time.time()

        segments_gen, info = self.model.transcribe(
            audio_path,
            language=self.config.language,
            beam_size=self.config.beam_size,
            best_of=self.config.best_of,
            patience=self.config.patience,
            vad_filter=self.config.vad_filter,
            vad_parameters=self.config.vad_parameters,
            word_timestamps=True,
            condition_on_previous_text=True,
            no_speech_threshold=0.6,
            log_prob_threshold=-1.0,
            compression_ratio_threshold=2.4,
            temperature=[0.0, 0.2, 0.4, 0.6, 0.8, 1.0],
            initial_prompt=INITIAL_PROMPT,
        )

        duration = info.duration
        console.print(
            f"[dim]Duration: {duration / 60:.1f} min | "
            f"Language: {info.language} ({info.language_probability:.0%})[/dim]"
        )

        raw = []
        with Progress(
            SpinnerColumn(), TextColumn("[progress.description]{task.description}"),
            BarColumn(), TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            TimeElapsedColumn(), console=console,
        ) as progress:
            task = progress.add_task("Processing...", total=duration)
            for seg in segments_gen:
                raw.append(Segment(start=seg.start, end=seg.end, text=seg.text.strip()))
                progress.update(task, completed=min(seg.end, duration))
            progress.update(task, completed=duration)

        elapsed = time.time() - start_time
        speed = duration / elapsed if elapsed > 0 else 0
        console.print(f"[green]✓[/green] Done in {elapsed:.1f}s ({speed:.1f}x realtime)")
        return raw

    def merge_segments(self, segments: list[Segment]) -> list[Segment]:
        """Merge fragmented whisper output into complete sentences."""
        if not segments:
            return []

        merged = []
        current = Segment(start=segments[0].start, end=segments[0].end, text=segments[0].text)

        for seg in segments[1:]:
            gap = seg.start - current.end
            text = current.text.strip()
            ends_sentence = text and text[-1] in ".!?"
            should_merge = (gap <= self.config.merge_gap_threshold and not ends_sentence) or gap < 0.5

            if should_merge:
                current.end = seg.end
                current.text = text + " " + seg.text.strip()
            else:
                if len(current.text.split()) >= self.config.min_segment_words:
                    merged.append(current)
                elif merged:
                    merged[-1].end = current.end
                    merged[-1].text += " " + current.text.strip()
                current = Segment(start=seg.start, end=seg.end, text=seg.text.strip())

        if current.text.strip():
            if len(current.text.split()) >= self.config.min_segment_words:
                merged.append(current)
            elif merged:
                merged[-1].end = current.end
                merged[-1].text += " " + current.text.strip()

        return merged

    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r'\b(\w+)( \1\b){2,}', r'\1', text)
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)
        text = re.sub(r'([.,!?;:])(\w)', r'\1 \2', text)
        text = re.sub(r'\s{2,}', ' ', text).strip()
        if text:
            text = text[0].upper() + text[1:]
            text = re.sub(r'(?<=[.!?]\s)(\w)', lambda m: m.group(1).upper(), text)
            if text[-1] not in '.!?':
                text += '.'
        return text

    def structure_by_flow(self, segments: list[Segment]) -> list[dict]:
        """
        Structure transcript into Q&A blocks using sentence-level analysis.

        Language-agnostic algorithm — works for any language using "?":
        1. Build timestamped sentence stream from Whisper segments
        2. Remove Whisper hallucinations (repeated content)
        3. Find question anchors: "?" sentences verified by audio gaps
        4. Handle command-style first questions (no "?")
        5. Build Q/A blocks, absorb short preambles
        6. Validate: merge tiny blocks, fix misclassified content
        """
        if not segments:
            return []

        sentences = self._build_sentences(segments)
        if not sentences:
            return []

        sentences = self._dedup_sentences(sentences)

        # ── Phase 1: Find question anchors ────────────────────────
        # A "?" sentence is a Q anchor if it's either:
        #   - Substantial (>= 10 words) — real question regardless of gap
        #   - Short but preceded by an audio gap (>= 0.4s) — speaker changed
        # Short "?" sentences with NO gap are likely rhetorical or quotes.
        gap_threshold = self.config.pause_for_turn_switch * 0.25
        q_anchors = []
        for i, s in enumerate(sentences):
            if "?" not in s["text"]:
                continue
            wc = len(s["text"].split())
            if wc < 4:
                continue
            gap = s["start"] - sentences[i - 1]["end"] if i > 0 else 999.0
            if wc >= 10 or gap >= gap_threshold:
                q_anchors.append(i)

        if not q_anchors:
            full = " ".join(s["text"] for s in sentences)
            return [{"type": "answer", "number": 0,
                     "text": self.clean_text(full),
                     "start": sentences[0]["start"],
                     "end": sentences[-1]["end"]}]

        # ── Phase 2: Group nearby anchors (same interviewer turn) ──
        # Two ? anchors belong in the same group ONLY IF:
        #   - They're within 3 sentences of each other AND
        #   - The content between them is short (< 20 words)
        # This prevents grouping Q-A-Q patterns where a short answer
        # sneaks in between two ? sentences.
        q_groups = []
        for idx in q_anchors:
            if q_groups:
                prev_idx = q_groups[-1][-1]
                gap = idx - prev_idx
                if gap <= 3:
                    between_words = sum(
                        len(sentences[j]["text"].split())
                        for j in range(prev_idx + 1, idx)
                        if "?" not in sentences[j]["text"]
                    )
                    if between_words < 20:
                        q_groups[-1].append(idx)
                        continue
            q_groups.append([idx])

        # ── Phase 3: Absorb preambles ─────────────────────────────
        # Look backward from each Q group's first sentence.
        # Absorb preceding sentence ONLY IF:
        #   - It's short (< 22 words, no "?")
        #   - There's a gap before IT (suggesting speaker change happened there)
        #   - Max 3 sentences / 50 words of preamble
        q_regions = []  # (start_idx, end_idx) inclusive
        for gi, group in enumerate(q_groups):
            q_end = group[-1]
            backward_limit = (q_regions[-1][1] + 1) if q_regions else 0
            q_start = group[0]
            preamble_words = 0

            for i in range(group[0] - 1, max(backward_limit - 1, -1), -1):
                wc = len(sentences[i]["text"].split())
                if wc > 22:
                    break  # too long — answer content
                if preamble_words + wc > 50:
                    break  # accumulated too much preamble
                # Check for gap before this sentence (speaker change signal)
                gap = sentences[i]["start"] - sentences[i - 1]["end"] if i > 0 else 999.0
                if gap < gap_threshold and i != 0:
                    break  # no gap = same speaker = still answer content
                q_start = i
                preamble_words += wc
                if group[0] - i >= 3:
                    break  # max 3 sentences of preamble

            q_regions.append((q_start, q_end))

        # ── Phase 4: First-block handling ─────────────────────────
        # If the transcript starts before the first Q anchor, detect
        # command-style questions ("Tell me about yourself." / no "?")
        first_q = q_regions[0][0]
        cmd_q_end = None

        if first_q >= 2:
            # First 1-3 short sentences could be a command question.
            # Answer starts where content becomes factual/long.
            q_word_total = 0
            for i in range(min(first_q, 4)):
                wc = len(sentences[i]["text"].split())
                if wc > 30:
                    break
                q_word_total += wc
                # Check: does the NEXT sentence look like answer content?
                if q_word_total >= 5 and i + 1 < first_q:
                    next_text = sentences[i + 1]["text"]
                    next_wc = len(next_text.split())
                    # Answer signals: starts with digit, or is a long
                    # declarative sentence (> 20 words, no "?")
                    is_digit = bool(re.match(r'\d', next_text))
                    is_long_decl = next_wc > 20 and "?" not in next_text
                    if is_digit or is_long_decl:
                        cmd_q_end = i + 1
                        break

        # ── Phase 5: Build blocks ─────────────────────────────────
        blocks = []
        q_num = 0

        # First block (command Q or pre-Q content)
        if cmd_q_end is not None:
            q_num += 1
            q_text = " ".join(sentences[i]["text"] for i in range(cmd_q_end))
            blocks.append({
                "type": "question", "number": q_num,
                "text": self.clean_text(q_text),
                "start": sentences[0]["start"],
                "end": sentences[cmd_q_end - 1]["end"],
            })
            if cmd_q_end < first_q:
                a_text = " ".join(
                    sentences[i]["text"] for i in range(cmd_q_end, first_q)
                )
                if a_text.strip():
                    blocks.append({
                        "type": "answer", "number": q_num,
                        "text": self.clean_text(a_text),
                        "start": sentences[cmd_q_end]["start"],
                        "end": sentences[first_q - 1]["end"],
                    })
        elif first_q > 0:
            # No command Q detected — dump as A0
            a_text = " ".join(sentences[i]["text"] for i in range(first_q))
            if a_text.strip():
                blocks.append({
                    "type": "answer", "number": 0,
                    "text": self.clean_text(a_text),
                    "start": sentences[0]["start"],
                    "end": sentences[first_q - 1]["end"],
                })

        # Main Q/A blocks
        for gi, (q_start, q_end) in enumerate(q_regions):
            q_num += 1
            q_text = " ".join(
                sentences[i]["text"] for i in range(q_start, q_end + 1)
            )
            blocks.append({
                "type": "question", "number": q_num,
                "text": self.clean_text(q_text),
                "start": sentences[q_start]["start"],
                "end": sentences[q_end]["end"],
            })

            a_start = q_end + 1
            a_end = q_regions[gi + 1][0] if gi + 1 < len(q_regions) else len(sentences)
            if a_start < a_end:
                a_text = " ".join(
                    sentences[i]["text"] for i in range(a_start, a_end)
                )
                if a_text.strip():
                    blocks.append({
                        "type": "answer", "number": q_num,
                        "text": self.clean_text(a_text),
                        "start": sentences[a_start]["start"],
                        "end": sentences[a_end - 1]["end"],
                    })

        # ── Phase 6: Validate ─────────────────────────────────────
        blocks = self._validate_blocks(blocks)

        # Renumber
        q_num = 0
        for b in blocks:
            if b["type"] == "question":
                q_num += 1
            b["number"] = q_num

        if self.guide_questions:
            idx = 0
            for b in blocks:
                if b["type"] == "question" and idx < len(self.guide_questions):
                    b["guide_question"] = self.guide_questions[idx]["text"]
                    idx += 1

        return blocks

    # ── Helpers ────────────────────────────────────────────────────

    def _build_sentences(self, segments: list[Segment]) -> list[dict]:
        """Split Whisper segments into individual timestamped sentences."""
        sentences = []
        for seg in segments:
            text = seg.text.strip()
            if not text:
                continue
            parts = re.split(r'(?<=[.!?])\s+', text)
            total_chars = sum(len(p) for p in parts) or 1
            duration = seg.end - seg.start
            offset = seg.start
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                d = (len(part) / total_chars) * duration
                sentences.append({
                    "text": part, "start": offset, "end": offset + d,
                    "seg_start": seg.start,
                })
                offset += d
        return sentences

    def _dedup_sentences(self, sentences: list[dict]) -> list[dict]:
        """Remove Whisper hallucinations (repeated content).

        Uses a sliding window to catch interleaved repeats, not just
        consecutive duplicates. Whisper sometimes repeats entire Q+A
        blocks like: [Q1][A1][Q1][A1][Q1]...
        """
        if len(sentences) < 2:
            return sentences
        window_size = 10
        result = [sentences[0]]
        for s in sentences[1:]:
            curr_w = set(s["text"].lower().split())
            if len(curr_w) <= 3:
                result.append(s)
                continue
            is_dup = False
            for prev in result[-window_size:]:
                prev_w = set(prev["text"].lower().split())
                if not prev_w:
                    continue
                overlap = len(prev_w & curr_w) / max(len(prev_w), len(curr_w))
                if overlap > 0.65:
                    sim = SequenceMatcher(
                        None, prev["text"].lower(), s["text"].lower()
                    ).ratio()
                    if sim > 0.55:
                        is_dup = True
                        break
            if not is_dup:
                result.append(s)
        return result

    def _validate_blocks(self, blocks: list[dict]) -> list[dict]:
        """
        Post-validation to fix structural errors.

        Rules:
        1. Q block > 50 words without "?" → reclassify as answer
        2. Q-A-Q where A < 12 words → merge (false split)
        3. Short Q (< 12w) that creates a short preceding A (< 25w) →
           likely rhetorical, merge back into answer
        4. Adjacent same-type blocks → merge
        """
        if not blocks:
            return blocks

        # Rule 1: Long Q without ? = misclassified answer
        for b in blocks:
            if b["type"] == "question" and "?" not in b["text"]:
                if len(b["text"].split()) > 50:
                    b["type"] = "answer"

        # Rule 2: Tiny answer between two Qs
        changed = True
        while changed:
            changed = False
            merged = []
            i = 0
            while i < len(blocks):
                if (
                    i + 2 < len(blocks)
                    and blocks[i]["type"] == "question"
                    and blocks[i + 1]["type"] == "answer"
                    and blocks[i + 2]["type"] == "question"
                    and len(blocks[i + 1]["text"].split()) < 12
                ):
                    blocks[i]["text"] += " " + blocks[i + 1]["text"] + " " + blocks[i + 2]["text"]
                    blocks[i]["end"] = blocks[i + 2]["end"]
                    merged.append(blocks[i])
                    i += 3
                    changed = True
                else:
                    merged.append(blocks[i])
                    i += 1
            blocks = merged

        # Rule 3: Short Q creating a short preceding A → rhetorical
        changed = True
        while changed:
            changed = False
            merged = []
            i = 0
            while i < len(blocks):
                if (
                    i >= 1
                    and blocks[i]["type"] == "question"
                    and len(blocks[i]["text"].split()) < 12
                    and blocks[i - 1]["type"] == "answer"
                    and len(blocks[i - 1]["text"].split()) < 25
                ):
                    # This Q is probably rhetorical — merge it and
                    # the short preceding A into the block before that
                    if i + 1 < len(blocks) and blocks[i + 1]["type"] == "answer":
                        # Merge: prev_A + this_Q + next_A → one big A
                        merged[-1]["text"] += (" " + blocks[i]["text"]
                                               + " " + blocks[i + 1]["text"])
                        merged[-1]["end"] = blocks[i + 1]["end"]
                        i += 2
                        changed = True
                    else:
                        merged[-1]["text"] += " " + blocks[i]["text"]
                        merged[-1]["end"] = blocks[i]["end"]
                        i += 1
                        changed = True
                else:
                    merged.append(blocks[i])
                    i += 1
            blocks = merged

        # Rule 4: Merge adjacent same-type blocks
        merged = [blocks[0]]
        for b in blocks[1:]:
            if b["type"] == merged[-1]["type"]:
                merged[-1]["text"] += " " + b["text"]
                merged[-1]["end"] = b["end"]
            else:
                merged.append(b)
        blocks = merged

        return blocks

    def process(self, audio_path: str) -> list[dict]:
        from rich.console import Console
        console = Console()

        raw = self.transcribe(audio_path)
        console.print(f"  Segments: {len(raw)} raw", end="")

        merged = self.merge_segments(raw)
        console.print(f" → {len(merged)} merged", end="")

        if self.guide_questions:
            blocks = self.structure_with_guide(merged)
            console.print(f" → {len(blocks)} Q&A blocks (guide-anchored)")
        else:
            blocks = self.structure_by_flow(merged)
            console.print(f" → {len(blocks)} Q&A blocks (flow-detected)")
        return blocks

    def structure_with_guide(self, segments: list[Segment]) -> list[dict]:
        """
        Use the interview guide to find Q/A boundaries via fuzzy matching.

        Strategy: for each guide question, extract key phrases and search
        segment-by-segment. Interviewers often rephrase, so we match on
        distinctive keywords rather than requiring exact sequences.
        """
        if not segments or not self.guide_questions:
            return []

        from rich.console import Console
        console = Console()

        def extract_keywords(text: str) -> set[str]:
            """Extract meaningful words (skip common Swedish/English stopwords)."""
            stops = {
                "och", "i", "på", "av", "en", "ett", "den", "det", "de", "som",
                "är", "var", "att", "till", "med", "för", "om", "kan", "har",
                "du", "dig", "din", "ditt", "dina", "vi", "jag", "så", "lite",
                "nu", "då", "ju", "hur", "vad", "när", "sig", "inte", "skulle",
                "the", "a", "an", "is", "are", "and", "or", "of", "to", "in",
                "you", "your", "do", "does", "can", "could", "would", "that",
                "this", "it", "be", "been", "have", "has", "was", "were",
            }
            words = re.findall(r'\w+', text.lower())
            return {w for w in words if len(w) > 2 and w not in stops}

        def match_score(guide_text: str, segment_text: str) -> float:
            """Score how well a segment matches a guide question."""
            guide_kw = extract_keywords(guide_text)
            seg_kw = extract_keywords(segment_text)

            if not guide_kw:
                return 0.0

            # Keyword overlap
            overlap = guide_kw & seg_kw
            kw_score = len(overlap) / len(guide_kw) if guide_kw else 0.0

            # Also check sequence similarity on first N words
            guide_start = " ".join(guide_text.lower().split()[:8])
            seg_start = " ".join(segment_text.lower().split()[:8])
            seq_score = SequenceMatcher(None, guide_start, seg_start).ratio()

            # Check if any distinctive multi-word phrases match
            guide_lower = guide_text.lower()
            seg_lower = segment_text.lower()
            phrase_bonus = 0.0

            # Extract 3-word phrases from guide
            guide_words = guide_lower.split()
            for j in range(len(guide_words) - 2):
                phrase = " ".join(guide_words[j:j+3])
                if phrase in seg_lower:
                    phrase_bonus = 0.3
                    break

            return max(kw_score, seq_score) + phrase_bonus

        def match_score_window(guide_text: str, seg_idx: int, window: int = 3) -> tuple[float, int, int]:
            """
            Try matching guide question against segments[seg_idx:seg_idx+window].
            Returns (best_score, start_seg, end_seg_exclusive).
            The question might span 1-3 segments.
            """
            best = 0.0
            best_range = (seg_idx, seg_idx + 1)

            for size in range(1, min(window + 1, len(segments) - seg_idx + 1)):
                combined = " ".join(segments[j].text for j in range(seg_idx, seg_idx + size))
                score = match_score(guide_text, combined)
                if score > best:
                    best = score
                    best_range = (seg_idx, seg_idx + size)

            return best, best_range[0], best_range[1]

        # Match each guide question to segments, searching forward only
        matches = []  # (guide_idx, seg_start, seg_end, score)
        search_from = 0

        for gi, gq in enumerate(self.guide_questions):
            best_score = 0.0
            best_start = -1
            best_end = -1

            # Search from current position forward
            for si in range(search_from, len(segments)):
                score, s_start, s_end = match_score_window(gq["text"], si)

                if score > best_score:
                    best_score = score
                    best_start = s_start
                    best_end = s_end

                # If we found a very good match, stop searching
                if best_score >= 0.6:
                    # But check a few more to be sure
                    for si2 in range(si + 1, min(si + 5, len(segments))):
                        score2, s2_start, s2_end = match_score_window(gq["text"], si2)
                        if score2 > best_score:
                            best_score = score2
                            best_start = s2_start
                            best_end = s2_end
                    break

            threshold = 0.25
            if best_score >= threshold and best_start >= 0:
                matches.append((gi, best_start, best_end, best_score))
                search_from = best_end
                console.print(
                    f"    [dim]Guide Q{gi+1}: seg {best_start}-{best_end-1} "
                    f"(score: {best_score:.2f}) "
                    f"\"{segments[best_start].text[:60]}...\"[/dim]"
                )
            else:
                console.print(f"    [dim]Guide Q{gi+1}: no match (best: {best_score:.2f})[/dim]")

        if not matches:
            console.print("  [yellow]No guide matches — falling back to flow detection[/yellow]")
            return self.structure_by_flow(segments)

        console.print(f"  [green]Matched {len(matches)}/{len(self.guide_questions)} guide questions[/green]")

        # Build Q&A blocks
        blocks = []

        # Handle any content before the first matched question
        if matches[0][1] > 0:
            pre_text = " ".join(segments[j].text for j in range(0, matches[0][1]))
            if pre_text.strip():
                blocks.append({
                    "type": "answer", "number": 0,
                    "text": self.clean_text(pre_text),
                    "start": segments[0].start,
                    "end": segments[matches[0][1] - 1].end,
                })

        for mi, (gi, q_start, q_end, score) in enumerate(matches):
            # Next question starts at...
            if mi + 1 < len(matches):
                next_q_start = matches[mi + 1][1]
            else:
                next_q_start = len(segments)

            # Question text
            q_text = " ".join(segments[j].text for j in range(q_start, q_end))
            blocks.append({
                "type": "question", "number": mi + 1,
                "text": self.clean_text(q_text),
                "start": segments[q_start].start,
                "end": segments[q_end - 1].end,
                "guide_question": self.guide_questions[gi]["text"],
            })

            # Answer = everything from q_end to next question
            if q_end < next_q_start:
                a_text = " ".join(segments[j].text for j in range(q_end, next_q_start))
                blocks.append({
                    "type": "answer", "number": mi + 1,
                    "text": self.clean_text(a_text),
                    "start": segments[q_end].start,
                    "end": segments[next_q_start - 1].end,
                })

        return blocks


# ──────────────────────────────────────────────────────────────
# Output Formatters
# ──────────────────────────────────────────────────────────────

def fmt_ts(seconds: float) -> str:
    h, remainder = divmod(int(seconds), 3600)
    m, s = divmod(remainder, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def save_docx(blocks: list[dict], path: str, source: str, guide: list[dict] = None):
    from docx import Document
    from docx.shared import Pt, Cm

    FONT = "Calibri"
    SIZE = Pt(11)

    def add(doc, text, bold=False, spacing_after=2):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(spacing_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = SIZE
        run.bold = bold
        return p

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Cm(2.5))

    add(doc, f"Interview Transcript — {source}", bold=True, spacing_after=4)
    add(doc, f"Transcribed: {datetime.now().strftime('%Y-%m-%d %H:%M')}", spacing_after=8)

    if guide:
        add(doc, "─" * 55)
        add(doc, "INTERVIEW GUIDE REFERENCE", bold=True, spacing_after=4)
        for q in guide:
            add(doc, f"  {q['number']}. {q['text']}")
        add(doc, "─" * 55, spacing_after=8)

    add(doc, "TRANSCRIPT", bold=True, spacing_after=8)

    for block in blocks:
        is_q = block["type"] == "question"
        label = f"{'Q' if is_q else 'A'}{block['number']}"
        ts = fmt_ts(block["start"])

        add(doc, f"{label} [{ts}]:", bold=True, spacing_after=2)
        add(doc, block["text"], spacing_after=8 if not is_q else 2)

    doc.save(path)


def save_txt(blocks: list[dict], path: str, source: str):
    lines = [f"INTERVIEW TRANSCRIPT — {source}", "=" * 60, ""]
    for block in blocks:
        prefix = f"{'Q' if block['type'] == 'question' else 'A'}{block['number']}"
        lines.append(f"{prefix} [{fmt_ts(block['start'])}]:")
        if "guide_question" in block:
            lines.append(f"  (Guide: {block['guide_question']})")
        lines.append(block["text"])
        lines.append("")
        if block["type"] == "answer":
            lines.extend(["-" * 40, ""])

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def save_json(blocks: list[dict], path: str, source: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"source": source, "blocks": blocks}, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Interview Transcription System",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
examples:
  python transcribe.py interview.mp3
  python transcribe.py interview.mp3 --guide questions.txt
  python transcribe.py interview.mp3 -o output.docx -f both
  python transcribe.py ./interviews/
  python transcribe.py interview.mp3 --language en --model medium
        """,
    )

    parser.add_argument("input", nargs="?", default=".", help="Audio file or directory (default: current directory)")
    parser.add_argument("--guide", "-g", help="Interview guide file")
    parser.add_argument("--output", "-o", help="Output path (auto-generated if omitted)")
    parser.add_argument("--format", "-f", choices=["docx", "txt", "both", "json"], default="docx")
    parser.add_argument("--model", "-m", default=MODEL_SIZE, help=f"Whisper model (default: {MODEL_SIZE})")
    parser.add_argument("--device", choices=["auto", "cuda", "cpu"], default="auto",
                        help="Device selection (default: auto)")
    parser.add_argument("--language", "-l", default=LANGUAGE, help=f"Language code (default: {LANGUAGE})")
    parser.add_argument("--beam-size", type=int, default=BEAM_SIZE, help=f"Beam size (default: {BEAM_SIZE})")
    parser.add_argument("--pause-threshold", type=float, default=PAUSE_THRESHOLD,
                        help=f"Speaker change pause in seconds (default: {PAUSE_THRESHOLD})")

    args = parser.parse_args()
    check_dependencies()

    from rich.console import Console
    from rich.panel import Panel
    console = Console()

    console.print(Panel.fit(
        f"[bold cyan]Interview Transcription System[/bold cyan]\n"
        f"Model: {args.model} | Device: {args.device} | Language: {args.language}",
        border_style="cyan",
    ))

    config = Config(
        model_size=args.model,
        device=args.device,
        language=args.language,
        beam_size=args.beam_size,
        guide_path=args.guide,
        output_format=args.format,
        pause_for_turn_switch=args.pause_threshold,
    )

    input_path = Path(args.input)
    if input_path.is_dir():
        audio_files = sorted(f for f in input_path.iterdir() if f.suffix.lower() in AUDIO_EXTENSIONS)
        if not audio_files:
            console.print(f"[red]No audio files found in {input_path}[/red]")
            sys.exit(1)
        console.print(f"Found [bold]{len(audio_files)}[/bold] audio files")
    elif input_path.is_file():
        audio_files = [input_path]
    else:
        console.print(f"[red]Not found: {args.input}[/red]")
        sys.exit(1)

    transcriber = Transcriber(config)

    for i, audio_file in enumerate(audio_files, 1):
        console.print(f"\n{'═' * 60}")
        console.print(f"[bold]File {i}/{len(audio_files)}:[/bold] {audio_file.name}")
        console.print(f"{'═' * 60}")

        blocks = transcriber.process(str(audio_file))

        if args.output and len(audio_files) == 1:
            out_base, out_dir = Path(args.output).stem, Path(args.output).parent
        else:
            out_base, out_dir = audio_file.stem + "_transcript", audio_file.parent

        if args.format in ("docx", "both"):
            p = out_dir / f"{out_base}.docx"
            save_docx(blocks, str(p), audio_file.name, transcriber.guide_questions)
            console.print(f"[green]✓[/green] {p}")

        if args.format in ("txt", "both"):
            p = out_dir / f"{out_base}.txt"
            save_txt(blocks, str(p), audio_file.name)
            console.print(f"[green]✓[/green] {p}")

        if args.format == "json":
            p = out_dir / f"{out_base}.json"
            save_json(blocks, str(p), audio_file.name)
            console.print(f"[green]✓[/green] {p}")

        console.print(f"\n[bold]Preview:[/bold]")
        for block in blocks[:6]:
            label = f"{'Q' if block['type'] == 'question' else 'A'}{block['number']}"
            style = "bold blue" if block["type"] == "question" else "bold green"
            text = block["text"][:150] + ("..." if len(block["text"]) > 150 else "")
            console.print(f"  [{style}]{label}[/{style}] [{fmt_ts(block['start'])}]: {text}")

    console.print(f"\n[bold green]Done.[/bold green] Processed {len(audio_files)} file(s).")


if __name__ == "__main__":
    main()